from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.config import settings
from app.converters.base import ConvertAsset, ConvertResult, rows_to_markdown_table


IMAGE_EXTENSIONS = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
    "image/webp": ".webp",
}


class DocxConverter:
    engine = "python-docx"

    def convert(self, input_path: Path, output_dir: Path) -> ConvertResult:
        document = Document(input_path)
        output_dir.mkdir(parents=True, exist_ok=True)
        assets_dir = output_dir / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)
        for old_asset in assets_dir.glob("image-*.*"):
            old_asset.unlink()

        parts: list[str] = []
        assets: list[ConvertAsset] = []
        image_index = 1

        for block in self._iter_body_blocks(document):
            if isinstance(block, Paragraph):
                markdown_blocks, image_index = self._paragraph_to_markdown(
                    block,
                    document,
                    output_dir,
                    assets_dir,
                    image_index,
                    assets,
                )
                parts.extend(markdown_blocks)
            elif isinstance(block, Table):
                table = self._table_to_markdown(block)
                if table:
                    parts.append(table)

        return ConvertResult(
            markdown="\n\n".join(parts),
            metadata={
                "engine": self.engine,
                "source_format": "docx",
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
                "asset_count": len(assets),
            },
            assets=assets,
        )

    def _iter_body_blocks(self, document: DocxDocument):
        for child in document.element.body.iterchildren():
            if not isinstance(child.tag, str):
                continue
            if child.tag.endswith("}p"):
                yield Paragraph(child, document)
            elif child.tag.endswith("}tbl"):
                yield Table(child, document)

    def _paragraph_to_markdown(
        self,
        paragraph: Paragraph,
        document: DocxDocument,
        output_dir: Path,
        assets_dir: Path,
        image_index: int,
        assets: list[ConvertAsset],
    ) -> tuple[list[str], int]:
        blocks: list[str] = []
        text = paragraph.text.strip()
        if text:
            blocks.append(self._format_paragraph_text(paragraph, text))

        document_id = output_dir.name
        asset_url = f"{settings.api_prefix}/documents/{document_id}/assets/"
        for image_part in self._iter_paragraph_images(paragraph, document):
            extension = self._image_extension(image_part.content_type, image_part.partname.ext)
            name = f"image-{image_index:03d}{extension}"
            path = assets_dir / name
            path.write_bytes(image_part.blob)
            blocks.append(f"![{name}]({asset_url}{name})")
            assets.append(
                ConvertAsset(
                    name=name,
                    content_type=image_part.content_type,
                    path=path.resolve(),
                )
            )
            image_index += 1

        return blocks, image_index

    def _iter_paragraph_images(self, paragraph: Paragraph, document: DocxDocument):
        for run in paragraph.runs:
            for blip in run._element.xpath(".//a:blip"):
                relationship_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                if relationship_id and relationship_id in document.part.related_parts:
                    yield document.part.related_parts[relationship_id]

    def _format_paragraph_text(self, paragraph: Paragraph, text: str) -> str:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading = re.match(r"Heading\s+([1-6])$", style_name or "")
        if heading:
            level = int(heading.group(1))
            return f"{'#' * level} {text}"
        if style_name == "Title":
            return f"# {text}"
        return text

    def _table_to_markdown(self, table: Table) -> str:
        rows: list[list[object]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        return rows_to_markdown_table(rows)

    def _image_extension(self, content_type: str | None, fallback_extension: str | None) -> str:
        if content_type in IMAGE_EXTENSIONS:
            return IMAGE_EXTENSIONS[content_type]
        if fallback_extension:
            extension = fallback_extension.lower()
            if re.fullmatch(r"\.[a-z0-9]{1,8}", extension):
                return extension
        return ".bin"
