from pathlib import Path
from uuid import uuid4

import pytest

from app.converters.base import ConversionError, write_convert_result
from app.converters.docling_converter import DoclingConverter
from app.converters.registry import get_converter


def work_dir(name: str) -> Path:
    path = Path(__file__).parent / "_tmp" / f"{name}_{uuid4().hex}"
    path.mkdir(parents=True, exist_ok=False)
    return path


def test_txt_converter() -> None:
    base_dir = work_dir("txt")
    source = base_dir / "demo.txt"
    source.write_text("hello\nworld", encoding="utf-8")

    result = get_converter("txt").convert(source, base_dir / "out")

    assert result.markdown == "hello\nworld"
    assert result.metadata["source_format"] == "txt"


def test_markdown_converter() -> None:
    base_dir = work_dir("markdown")
    source = base_dir / "demo.md"
    source.write_text("# Title", encoding="utf-8")

    result = get_converter("markdown").convert(source, base_dir / "out")

    assert result.markdown == "# Title"
    assert result.metadata["source_format"] == "markdown"


def test_csv_converter() -> None:
    base_dir = work_dir("csv")
    source = base_dir / "demo.csv"
    source.write_text("name,age\nAlice,18\nBob,20", encoding="utf-8")

    result = get_converter("csv").convert(source, base_dir / "out")

    assert "| name | age |" in result.markdown
    assert "| Alice | 18 |" in result.markdown
    assert result.metadata["rows"] == 3


def test_xlsx_converter() -> None:
    from openpyxl import Workbook

    base_dir = work_dir("xlsx")
    source = base_dir / "demo.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Users"
    sheet.append(["name", "age"])
    sheet.append(["Alice", 18])
    workbook.save(source)

    result = get_converter("xlsx").convert(source, base_dir / "out")

    assert "## Users" in result.markdown
    assert "| name | age |" in result.markdown
    assert result.metadata["sheets"][0]["name"] == "Users"


def test_html_converter() -> None:
    base_dir = work_dir("html")
    source = base_dir / "demo.html"
    source.write_text(
        "<html><head><title>Demo</title></head><body>"
        "<h1>Title</h1><p>Paragraph</p><ul><li>One</li></ul>"
        "</body></html>",
        encoding="utf-8",
    )

    result = get_converter("html").convert(source, base_dir / "out")

    assert "# Title" in result.markdown
    assert "Paragraph" in result.markdown
    assert "- One" in result.markdown
    assert result.metadata["title"] == "Demo"


def test_write_convert_result() -> None:
    base_dir = work_dir("write")
    source = base_dir / "demo.txt"
    source.write_text("hello", encoding="utf-8")
    result = get_converter("txt").convert(source, base_dir / "out")

    markdown_path, metadata_path = write_convert_result(result, base_dir / "out")

    assert markdown_path.read_text(encoding="utf-8") == "hello"
    assert metadata_path.exists()


def test_registry_docling_formats() -> None:
    assert get_converter("pdf").engine == "docling"
    assert get_converter("docx").engine == "docling"
    assert get_converter("pptx").engine == "docling"


def test_docling_docx_fallback_exports_tables_and_images(monkeypatch) -> None:
    from io import BytesIO

    from docx import Document
    from lxml import etree
    from PIL import Image

    class FailingDoclingConverter:
        def convert(self, path: str) -> None:
            raise RuntimeError("docling failed")

    monkeypatch.setattr(
        DoclingConverter,
        "_create_converter",
        lambda self: FailingDoclingConverter(),
    )

    base_dir = work_dir("docx")
    source = base_dir / "demo.docx"
    document = Document()
    document.element.body.insert(0, etree.Comment("comment before body content"))
    document.add_heading("Docx Title", level=1)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "age"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "18"
    image = BytesIO()
    Image.new("RGB", (16, 16), color="red").save(image, format="PNG")
    image.seek(0)
    document.add_picture(image)
    document.save(source)

    result = get_converter("docx").convert(source, base_dir / "out" / "demo-file-id")

    assert "# Docx Title" in result.markdown
    assert "| name | age |" in result.markdown
    assert "Alice" in result.markdown
    assert "![image-001.png](/api/documents/demo-file-id/assets/image-001.png)" in result.markdown
    assert result.assets[0].name == "image-001.png"
    assert result.assets[0].path.exists()
    assert result.metadata["fallback_engine"] == "python-docx"
    assert result.warnings == ["Docling处理DOCX失败，已使用python-docx兜底转换"]


def test_docling_dependency_available() -> None:
    import docling

    assert docling is not None


def test_registry_unsupported_format() -> None:
    with pytest.raises(ConversionError) as exc:
        get_converter("png")

    assert exc.value.error_code == "unsupported_file_format"
